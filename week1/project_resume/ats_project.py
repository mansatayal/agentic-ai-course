import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
import json

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("api key not found")

client = Groq(api_key = my_api_key)
model = "openai/gpt-oss-120b"


# Job Description schema -------------------------------------------------------------

job_description = """
About the job
Key responsibilities:

1. Develop and deploy AI-powered solutions and applications.
2. Work with LLMs such as GPT, Claude, Gemini, and open-source models.
3. Research and implement the latest AI technologies and tools.
4. Integrate AI models and APIs into scalable software applications.
5. Collaborate with the development team to test, optimize, and improve AI solutions.
Skill(s) required
Artificial intelligence
Git
GitHub
LLMOps
Machine Learning
Python
REST API
Earn certifications in these skills
Learn Artificial intelligence
Learn Git
Learn LLMOps
Learn Machine Learning
Learn Python
Other requirements
1. Candidates with strong Python and AI/ML fundamentals.

2. Hands-on experience with Generative AI, LLMs, or AI projects.

3. Strong problem-solving skills and willingness to learn new technologies.

4. Candidates with good communication skills and the ability to work effectively in a team.

Salary
Annual CTC: ₹ 2,40,000 - 3,00,000 /year

Perks
Informal dress code
Number of openings
4
About Stirring Minds
Stirring Minds is a premier startup ecosystem in India, dedicated to helping businesses launch, scale, and succeed. As a leading incubator, we provide funding, co-working spaces, and mentorship to support the growth of innovative companies. In addition to our incubator services, we also host the largest startup event in the country known as Startup Summit Live, bringing together entrepreneurs and industry leaders to connect, learn, and collaborate. Our community-driven approach extends beyond our event and incubator offerings, as we work to create communities of like-minded individuals who can support and learn from one another. We have been recognized by top media outlets both in India and internationally, including the BBC, The Guardian, Entrepreneur, and Business Insider. Our goal is to provide a comprehensive ecosystem for startups and help turn their ideas into reality.

"""
class  jobD(BaseModel):
    required_skills : list[str]
    preffered_skills : list[str]
    location : str
    minimum_experience : int | None         #makes it optional
    minimum_qualifications : list[str]
    salary : int
    responsibilities : list[str]

jobD_schema = jobD.model_json_schema()

system_prompt = f"""
You're a senior HR. Your job is to analyze job description and extract structured information from them.
Return only valid JSON matching this schema {jobD_schema}

IMPORTANT:
Do not return the schema itself and fill the schema with actual information extracted form the job description.
If minimum experience is not mentioned return null.
If information for a list is missing return an empty list 
Do not invent information.

"""

user_prompt = f"Analyze the job description {job_description}"

message_system = {
    "role" : "system",
    "content" : system_prompt
} 

message_user = {
    "role" : "user",
    "content" : user_prompt
}

response_format = {
    "type" : "json_object"
}

messages = [message_system, message_user]

response = client.chat.completions.create(model = model, messages = messages, response_format = response_format)

answer = response.choices[0].message.content

raw_json = answer
data_file = json.loads(raw_json)
jd = jobD(**data_file)

print(f"""   
THIS IS WHAT THE JOB DESCRIPTION SAYS:
required_skills : {jd.required_skills}
preffered_skills : {jd.preffered_skills}
location : {jd.location}
minimum_experience : {jd.minimum_experience}
minimum_qualifications : {jd.minimum_qualifications}
salary : {jd.salary}
responsibilities : {jd.responsibilities}
""")




# resume schema -------------------------------------------------------------

class experience(BaseModel):
    company : str | None
    role : str| None
    duration: str | None
    description : str | None
    skills_used : list[str] = []

class resume(BaseModel):
    name : str | None = None
    email : str | None  = None
    phone : str | None = None       #string because numbers may have leading zeroes and country code(+91 ) with a space that int would reject
    total_experience_years : float | None = None
    skills : list[str] = []
    experinces : list[experience] = []
    education : list[str] = []
    projects : list[str] = []
    certificates : list[str] = []

resume_schema = resume.model_json_schema()



# score -------------------------------------------------------------------------
class match_result(BaseModel):
    score : float
    details : dict


def final_score(job, resume):
    match_schema = match_result.model_json_schema()

    prompt = f"""
    You're a senior HR recruiter compare the candidate's resume with the job description.
    job description :  {job.model_dump_json(indent = 2)}
    candidate resume : {resume.model_dump_json(indent = 2)}
    Return JSON matching this schema : {match_schema}

    Give me:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.

    """

    message = {
        "role" : "user",
        "content" : prompt 
    }

    messages = [message]
    response_format = {"type" : "json_object"}

    response = client.chat.completions.create(model = model, messages = messages, response_format = response_format)

    data = json.loads(response.choices[0].message.content)
    return match_result(**data)


# read resume ---------------------------------------------------------
def parse_resume(resume_text):
    system_prompt = f"""
    You're an expert resume parser. Extract information from the resume from it's meaning not just it's heading.
    
    Different resumes may use different headings for example: Experince, Professional Experince, Work History, Empluments, Internships etc these may all contain relevant experince.

    Return only valid JSON matching this schema: {resume_schema}

    IMPORTANT:
    Don't invent information
    If a value is not available return NULL
    If a list has no information return empty list
    Extract skills mentioned across the resume

    """

    user_prompt = f"Parse this resume : {resume_text}"

    message_system = {
        "role" : "system",
        "content" : system_prompt
    }

    message_user = {
        "role" : "user",
        "content" : user_prompt
    }

    messages = [message_system, message_user]

    response_format = {"type" : "json_object"}
    response = client.chat.completions.create(model = model, messages = messages, response_format = response_format)

    raw_json = response.choices[0].message.content
    data = json.loads(raw_json)
    parsed = resume(**data)
    return parsed




# read from pdf or word file-------------------------------------------------------------
from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def read_docx(file_path):
    document = Document(file_path)
    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None




# read resume files from the folder "resume"-----------------------------------------
resume_folder = Path("resumes")
all_results = []

for file_path in resume_folder.iterdir():   #for files in resume folder
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print("\nProcessing: ", file_path.name)

    resume_text = read_resume(file_path)

    parsed_resume = parse_resume(resume_text)   #first llm call
    time.sleep(5)       #wait for 5 secs so that if there are multiple files so groq don't block them

    result = final_score(jd, parsed_resume)    #second llm call
    time.sleep(5)
    print("score: ", result.score)

    all_results.append({
        "name" : parsed_resume.name,
        "score" : result.score,
        "details" : result.details
    })



all_results.sort(
    key = lambda candidate: candidate["score"],     #it's a quick one line function
    reverse=True                                    #sort best to worst
)

top_2 = all_results[:2]
worst_2 = all_results[-2:]

print("\nTop two candidates:")
for candidate in top_2:
    print(candidate["name"], "-" , candidate["score"], "%")
    print(candidate["details"])

print("\nLowest Two candidates:")
for candidate in worst_2:
    print(candidate["name"], "-" , candidate["score"], "%")
    print(candidate["details"])
